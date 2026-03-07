from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

DOC_PATH = "C:\\Users\\root\\Desktop\\刘伟\\刘伟-feedback3.docx"
CAPTION = "表3-2 商家表"
HEADER = ["字段名", "数据类型", "主键/外键", "约束", "字段说明"]
DATA = [
    ["id", "int(11)", "主键", "自增、非空", "商家唯一ID"],
    ["phone", "varchar(11)", "无", "非空、唯一", "商家手机号（登录账号）"],
    ["password", "varchar(100)", "无", "非空", "加密后的登录密码"],
    ["shop_name", "varchar(100)", "无", "非空", "店铺名称"],
    ["license", "varchar(255)", "无", "非空", "营业执照URL"],
    ["delivery_range", "varchar(255)", "无", "非空", "配送范围描述"],
    ["status", "int(2)", "无", "非空", "状态（0-待审核，1-营业，2-禁用）"],
]


def set_cell_border(cell, **edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in('w:tcBorders')
    if tc_borders is None:
        tc_borders = OxmlElement('w:tcBorders')
        tc_pr.append(tc_borders)
    for edge_name, edge_data in edges.items():
        tag = 'w:' + edge_name
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in edge_data.items():
            element.set(qn('w:' + key), str(value))


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.first_child_found_in('w:tblBorders')
    if tbl_borders is None:
        tbl_borders = OxmlElement('w:tblBorders')
        tbl_pr.append(tbl_borders)
    desired = {
        'top': {'val': 'single', 'sz': 12, 'space': 0, 'color': '000000'},
        'left': {'val': 'nil'},
        'bottom': {'val': 'single', 'sz': 12, 'space': 0, 'color': '000000'},
        'right': {'val': 'nil'},
        'insideH': {'val': 'nil'},
        'insideV': {'val': 'nil'},
    }
    for edge_name, edge_data in desired.items():
        tag = qn('w:' + edge_name)
        element = tbl_borders.find(tag)
        if element is None:
            element = OxmlElement('w:' + edge_name)
            tbl_borders.append(element)
        for key, value in edge_data.items():
            element.set(qn('w:' + key), str(value))


def format_cell(cell, text, align=WD_ALIGN_PARAGRAPH.CENTER, bold=False):
    cell.text = text
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = align
        for run in paragraph.runs:
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(10.5)
            run.bold = bold


def remove_row(table, row_idx):
    tr = table.rows[row_idx]._tr
    tr.getparent().remove(tr)


def find_target_table(doc):
    body = doc._element.body
    found_caption = False
    table_ordinal = -1
    for child in body.iterchildren():
        tag = child.tag.split('}')[-1]
        if tag == 'p':
            text = ''.join(child.xpath('.//w:t/text()')).strip()
            if text == CAPTION:
                found_caption = True
                continue
        elif tag == 'tbl':
            table_ordinal += 1
            if found_caption:
                return doc.tables[table_ordinal]
    return None


doc = Document(DOC_PATH)
table = find_target_table(doc)
if table is None:
    raise SystemExit('ERROR: 未找到表3-2 商家表')

for idx in range(len(table.rows) - 1, 0, -1):
    values = [cell.text.strip() for cell in table.rows[idx].cells]
    if values == HEADER:
        remove_row(table, idx)

expected_rows = 1 + len(DATA)
while len(table.rows) < expected_rows:
    table.add_row()
while len(table.rows) > expected_rows:
    remove_row(table, len(table.rows) - 1)

for col_idx, value in enumerate(HEADER):
    format_cell(table.rows[0].cells[col_idx], value, bold=True)

for row_idx, row_data in enumerate(DATA, start=1):
    for col_idx, value in enumerate(row_data):
        align = WD_ALIGN_PARAGRAPH.LEFT if col_idx == 4 else WD_ALIGN_PARAGRAPH.CENTER
        format_cell(table.rows[row_idx].cells[col_idx], value, align=align)

table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(table)
for cell in table.rows[0].cells:
    set_cell_border(cell, bottom={'val': 'single', 'sz': 4, 'space': 0, 'color': '000000'})
for row in table.rows[1:]:
    for cell in row.cells:
        set_cell_border(
            cell,
            top={'val': 'nil'},
            left={'val': 'nil'},
            bottom={'val': 'nil'},
            right={'val': 'nil'},
        )

doc.save(DOC_PATH)

# read back and verify
verified = Document(DOC_PATH)
verified_table = find_target_table(verified)
print('已更新:', DOC_PATH)
print('行数:', len(verified_table.rows), '列数:', len(verified_table.columns))
for idx, row in enumerate(verified_table.rows):
    print(idx, [cell.text.strip() for cell in row.cells])
